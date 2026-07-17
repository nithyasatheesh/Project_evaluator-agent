"""
readers.py
==========
Format-specific content readers. Each reader takes a Path to a file
already extracted on disk and returns a plain-text representation of
its content suitable for LLM consumption.

All readers raise ReaderError on failure; callers (parser.py) are
expected to catch ReaderError per-file and continue processing the
rest of the submission rather than aborting.
"""

from __future__ import annotations

import json
from pathlib import Path

import config
from utils import safe_decode, setup_logger, truncate_text

logger = setup_logger(__name__)


class ReaderError(Exception):
    """Raised when a file cannot be read or parsed by its format reader."""


def _read_bytes_limited(path: Path) -> bytes:
    """Read up to MAX_FILE_READ_BYTES from a file."""
    try:
        with path.open("rb") as fh:
            return fh.read(config.MAX_FILE_READ_BYTES)
    except OSError as exc:
        raise ReaderError(f"Could not open file '{path.name}': {exc}") from exc


def read_text_file(path: Path) -> str:
    """Read a plain-text file (.txt, .md, .py, .sh, .sql, .json, .yaml, .yml)."""
    raw = _read_bytes_limited(path)
    try:
        return safe_decode(raw)
    except Exception as exc:  # pragma: no cover - safe_decode already very tolerant
        raise ReaderError(f"Could not decode text file '{path.name}': {exc}") from exc


def read_pdf(path: Path) -> str:
    """Extract text content from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ReaderError("pypdf is not installed") from exc

    try:
        reader = PdfReader(str(path))
        pages_text = []
        for page in reader.pages:
            try:
                pages_text.append(page.extract_text() or "")
            except Exception as page_exc:  # noqa: BLE001
                logger.warning("Failed to extract a page from PDF '%s': %s", path.name, page_exc)
        text = "\n".join(pages_text).strip()
        if not text:
            raise ReaderError(f"PDF '{path.name}' contains no extractable text (may be scanned/image-based).")
        return text
    except ReaderError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ReaderError(f"Could not parse PDF '{path.name}': {exc}") from exc


def read_docx(path: Path) -> str:
    """Extract paragraph and table text from a Word document."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ReaderError("python-docx is not installed") from exc

    try:
        document = docx.Document(str(path))
        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        raise ReaderError(f"Could not parse DOCX '{path.name}': {exc}") from exc


def read_html(path: Path) -> str:
    """Strip tags from an HTML file and return visible text content."""
    raw = _read_bytes_limited(path)
    try:
        from bs4 import BeautifulSoup

        text = safe_decode(raw)
        soup = BeautifulSoup(text, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        content = soup.get_text(separator="\n")
        lines = [line.strip() for line in content.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception as exc:  # noqa: BLE001
        raise ReaderError(f"Could not parse HTML '{path.name}': {exc}") from exc


def read_ipynb(path: Path) -> str:
    """Extract markdown, code and text output cells from a Jupyter notebook."""
    raw = _read_bytes_limited(path)
    try:
        notebook = json.loads(safe_decode(raw))
    except json.JSONDecodeError as exc:
        raise ReaderError(f"Notebook '{path.name}' is not valid JSON: {exc}") from exc

    parts: list[str] = []
    for cell in notebook.get("cells", []):
        cell_type = cell.get("cell_type", "")
        source = cell.get("source", "")
        source_text = "".join(source) if isinstance(source, list) else str(source)
        if not source_text.strip():
            continue
        if cell_type == "markdown":
            parts.append(f"[MARKDOWN]\n{source_text}")
        elif cell_type == "code":
            parts.append(f"[CODE]\n{source_text}")
            for output in cell.get("outputs", []) or []:
                text_out = output.get("text")
                if text_out:
                    out_str = "".join(text_out) if isinstance(text_out, list) else str(text_out)
                    parts.append(f"[OUTPUT]\n{out_str}")
                data = output.get("data", {})
                if isinstance(data, dict) and "text/plain" in data:
                    plain = data["text/plain"]
                    plain_str = "".join(plain) if isinstance(plain, list) else str(plain)
                    parts.append(f"[OUTPUT]\n{plain_str}")
    if not parts:
        raise ReaderError(f"Notebook '{path.name}' contains no readable cell content.")
    return "\n\n".join(parts)


def read_csv_summary(path: Path) -> str:
    """Produce a compact textual summary of a CSV file (shape, columns, preview)."""
    try:
        import pandas as pd
    except ImportError as exc:  # pragma: no cover
        raise ReaderError("pandas is not installed") from exc

    try:
        df = pd.read_csv(path, nrows=config.MAX_CSV_PREVIEW_ROWS, on_bad_lines="skip")
        try:
            total_rows = sum(1 for _ in open(path, "r", encoding="utf-8", errors="replace")) - 1
        except OSError:
            total_rows = len(df)
        summary_lines = [
            f"File: {path.name}",
            f"Approx rows: {max(total_rows, len(df))}, Columns: {len(df.columns)}",
            f"Column names: {', '.join(str(c) for c in df.columns)}",
            "Preview:",
            df.head(config.MAX_CSV_PREVIEW_ROWS).to_string(index=False),
        ]
        return "\n".join(summary_lines)
    except Exception as exc:  # noqa: BLE001
        raise ReaderError(f"Could not parse CSV '{path.name}': {exc}") from exc


def read_config_file(path: Path) -> str:
    """Read a JSON/YAML/YML configuration file as text (with validation)."""
    raw_text = read_text_file(path)
    suffix = path.suffix.lower()
    try:
        if suffix == ".json":
            json.loads(raw_text)  # validate only; keep original formatting
        elif suffix in (".yaml", ".yml"):
            import yaml

            yaml.safe_load(raw_text)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Config file '%s' failed validation (kept as raw text): %s", path.name, exc)
    return raw_text


def describe_image(path: Path) -> str:
    """Return a short descriptor string for an image file (name + size)."""
    try:
        size_bytes = path.stat().st_size
        return f"{path.name} ({size_bytes / 1024:.1f} KB)"
    except OSError as exc:
        raise ReaderError(f"Could not stat image '{path.name}': {exc}") from exc


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------
def read_file_content(path: Path) -> str:
    """Dispatch to the correct reader based on file extension and truncate.

    Raises ReaderError if the extension is unsupported or the underlying
    reader fails.
    """
    suffix = path.suffix.lower()

    if suffix in config.DOC_EXTENSIONS:
        if suffix == ".pdf":
            content = read_pdf(path)
        elif suffix == ".docx":
            content = read_docx(path)
        else:
            content = read_text_file(path)
    elif suffix in config.NOTEBOOK_EXTENSIONS:
        content = read_ipynb(path)
    elif suffix in config.HTML_EXTENSIONS:
        content = read_html(path)
    elif suffix in config.PYTHON_EXTENSIONS or suffix in config.SHELL_EXTENSIONS or suffix in config.SQL_EXTENSIONS:
        content = read_text_file(path)
    elif suffix in config.CSV_EXTENSIONS:
        content = read_csv_summary(path)
    elif suffix in config.CONFIG_EXTENSIONS:
        content = read_config_file(path)
    else:
        raise ReaderError(f"Unsupported file extension '{suffix}' for '{path.name}'")

    return truncate_text(content)
